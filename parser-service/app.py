import io
from fastapi import FastAPI, UploadFile, File, Body
from fastapi.middleware.cors import CORSMiddleware
from parser import pdf_parser, docx_parser
import os, tempfile
from groq import Groq
from PyPDF2 import PdfReader
from docx import Document
import re
import nltk
from nltk import pos_tag
from nltk.corpus import stopwords
from fastapi import FastAPI, Form
from fastapi.responses import JSONResponse
from sentence_transformers import SentenceTransformer, util


GROQ_API_KEY='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'
client = Groq(api_key=GROQ_API_KEY)
model = SentenceTransformer('BAAI/bge-large-en')
nltk.download('punkt_tab')
nltk.download('averaged_perceptron_tagger')
nltk.download('stopwords')
nltk.download('averaged_perceptron_tagger_eng')
from nltk.tokenize import sent_tokenize
app = FastAPI()

# CORS (allow all for dev)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------
# Base Resume Parser
# ---------------------------
@app.post("/parse")
async def parse_resume(resume: UploadFile = File(...)):
    suffix = os.path.splitext(resume.filename)[1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await resume.read())
        tmp_path = tmp.name

    try:
        if suffix == ".pdf":
            parsed = pdf_parser.parse(tmp_path)
        elif suffix == ".docx":
            parsed = docx_parser.parse(tmp_path)
        else:
            parsed = {"error": "Unsupported file format"}
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

    return parsed


# ---------------------------
# ATS Tester
# ---------------------------
@app.post("/ats-test")
# async def ats_test(resume: UploadFile = File(...), job_desc: str = Form(...)):
#     # Step 1: Parse resume
#     contents = await resume.read()
#     suffix = os.path.splitext(resume.filename)[1].lower()
#     if suffix == ".pdf":
#         from PyPDF2 import PdfReader
#         reader = PdfReader(io.BytesIO(contents))
#         resume_text = "".join([page.extract_text() or "" for page in reader.pages]).lower()
#     elif suffix == ".docx":
#         from docx import Document
#         doc = Document(io.BytesIO(contents))
#         resume_text = "\n".join([para.text for para in doc.paragraphs]).lower()
#     else:
#         return {"error": "Unsupported file format"}


# async def ats_test(resumes: list[UploadFile] = File(...), job_desc: str = Form(...)):
#     results = []

#     # ensure job description lowercase
#     job_desc = job_desc.lower()

#     for resume in resumes:
#         # read resume content
#         contents = await resume.read()
#         suffix = os.path.splitext(resume.filename)[1].lower()

#         # extract text
#         try:
#             if suffix == ".pdf":
#                 reader = PdfReader(io.BytesIO(contents))
#                 resume_text = "".join([page.extract_text() or "" for page in reader.pages]).lower()
#             elif suffix == ".docx":
#                 doc = Document(io.BytesIO(contents))
#                 resume_text = "\n".join([para.text for para in doc.paragraphs]).lower()
#             else:
#                 results.append({
#                     "filename": resume.filename,
#                     "error": "Unsupported file format (use .pdf or .docx)"
#                 })
#                 continue
#         except Exception as e:
#             results.append({
#                 "filename": resume.filename,
#                 "error": f"Error reading file: {str(e)}"
#             })
#             continue
# # Step 2: Prepare job tokens
#     # --------------------------
#     stop_words = {
#     "a", "an", "the", "in", "on", "and", "of", "to", "for", "from", "with",
#     "at", "by", "this", "that", "is", "it", "as", "are", "was", "be", "or"
#     }
    
#     job_tokens = {w for w in re.findall(r"\b\w+\b", job_desc.lower()) if w not in stop_words and len(w) > 2}
#     resume_tokens = {w for w in re.findall(r"\b\w+\b", resume_text.lower()) if w not in stop_words and len(w) > 2}
#     matched = resume_tokens & job_tokens


#     # Step 2: Section detection
#     ats_sections = {
#     "summary": ["summary", "profile", "overview", "professional summary"],
#     "objective": ["objective", "career objective"],
#     "experience": ["experience", "work experience", "employment", "professional experience"],
#     "education": ["education", "qualifications", "academics", "academic background"],
#     "skills": ["skills", "technical skills", "expertise", "competencies"],
#     "projects": ["projects", "work samples", "portfolio", "personal projects"],
#     "certifications": ["certifications", "licenses", "achievements", "awards"]
#     }

#     sections_detected = {sec: any(k.lower() in resume_text for k in kws) for sec, kws in ats_sections.items()}


#   # Step 4: Semantic similarity using Transformer
#     resume_sents = sent_tokenize(resume_text)
#     jd_sents = sent_tokenize(job_desc)

#     resume_embeddings = model.encode(resume_sents, convert_to_tensor=True)
#     jd_embeddings = model.encode(jd_sents, convert_to_tensor=True)

# # Compute cosine similarity between each JD sentence and each resume sentence
#     cos_scores = util.pytorch_cos_sim(jd_embeddings, resume_embeddings)

# # Take the max similarity for each JD sentence, then average
#     max_sim_per_jd = cos_scores.max(dim=1).values
#     semantic_similarity = max_sim_per_jd.mean().item()

#     # Step 5: Compute scores
#     section_score = sum(sections_detected.values()) * 5
#     keyword_score = len(matched) / max(len(job_tokens), 1) * 35
#     semantic_score = semantic_similarity * 60  # scaled to 0-50
#     if semantic_similarity < 0.6:
#        semantic_score *= 0.5
#     ats_score = round(min(100, section_score + keyword_score + semantic_score), 2)

#     # Step 6: Suggestions
#     suggestions = []

#     job_keywords = list(job_tokens - resume_tokens)
#     if job_keywords:
#         # Get embeddings for job description keywords and full resume
#         job_kw_embeddings = model.encode(job_keywords, convert_to_tensor=True)
#         resume_embedding = model.encode(resume_text, convert_to_tensor=True)

#         # Compute cosine similarity between each job keyword and entire resume
#         sims = util.cos_sim(job_kw_embeddings, resume_embedding).squeeze()

#         # Rank by *lowest* semantic similarity (i.e., most missing concepts)
#         ranked_keywords = sorted(
#             zip(job_keywords, sims.tolist()),
#             key=lambda x: x[1]
#         )

#         # Filter out very common or unimportant words
#         ignore = {"high", "good", "able", "using", "based", "help", "work", "team",
#           "time", "will", "you", "your", "well", "effort", "performing",
#           "leading", "cross", "reviewing", "responsible", "ensure", "support"}
      

#         # Extract only nouns, verbs, and adjectives (technical/meaningful words)
#         filtered_keywords = []
#         for kw, score in ranked_keywords:
#            if kw in ignore or kw in stopwords.words('english'):
#                 continue
#            pos = pos_tag([kw])[0][1]
#            if pos.startswith(('NN', 'VB', 'JJ')):  # Noun, Verb, Adjective
#                 filtered_keywords.append(kw)

#         top_keywords = filtered_keywords[:17]  # top 7 missing keywords

#         if top_keywords:
#             suggestions.append(
#                f"Consider adding these relevant keywords to better align with the job description: {', '.join(top_keywords)}."
#             )
#         else:
#             suggestions.append("Your resume semantically covers most of the job description well.")
#     else:
#         suggestions.append("No missing keywords detected.")

#     results.append ({
#         "ats_score": ats_score,
#         "sections_detected": sections_detected,
#         "semantic_similarity": round(semantic_similarity * 100, 2),
#         "keyword_match": {
#             "match_percent": round(len(matched) / max(len(job_tokens), 1) * 100, 2),
#             "matched_keywords": sorted(matched),
#         },
#         "improvements": suggestions,
#     })
#     if len(results) == 1:
#       return results[0]
#     else:
#       return {"results": results}


@app.post("/api/ats-test")
async def ats_test(resumes: list[UploadFile] = File(...), job_desc: str = Form(...)):
    results = []

    # ensure job description lowercase
    job_desc = job_desc.lower()

    for resume in resumes:
        # Step 1: Read resume content
        contents = await resume.read()
        suffix = os.path.splitext(resume.filename)[1].lower()

        # Extract text
        try:
            if suffix == ".pdf":
                reader = PdfReader(io.BytesIO(contents))
                resume_text = "".join([page.extract_text() or "" for page in reader.pages]).lower()
            elif suffix == ".docx":
                doc = Document(io.BytesIO(contents))
                resume_text = "\n".join([para.text for para in doc.paragraphs]).lower()
            else:
                results.append({
                    "filename": resume.filename,
                    "error": "Unsupported file format (use .pdf or .docx)"
                })
                continue
        except Exception as e:
            results.append({
                "filename": resume.filename,
                "error": f"Error reading file: {str(e)}"
            })
            continue

        # Step 2: Prepare job tokens
        stop_words = {
            "a", "an", "the", "in", "on", "and", "of", "to", "for", "from", "with",
            "at", "by", "this", "that", "is", "it", "as", "are", "was", "be", "or"
        }

        job_tokens = {w for w in re.findall(r"\b\w+\b", job_desc) if w not in stop_words and len(w) > 2}
        resume_tokens = {w for w in re.findall(r"\b\w+\b", resume_text) if w not in stop_words and len(w) > 2}
        matched = resume_tokens & job_tokens

        # Step 3: Section detection
        ats_sections = {
            "summary": ["summary", "profile", "overview", "professional summary"],
            "objective": ["objective", "career objective"],
            "experience": ["experience", "work experience", "employment", "professional experience"],
            "education": ["education", "qualifications", "academics", "academic background"],
            "skills": ["skills", "technical skills", "expertise", "competencies"],
            "projects": ["projects", "work samples", "portfolio", "personal projects"],
            "certifications": ["certifications", "licenses", "achievements", "awards"]
        }

        sections_detected = {sec: any(k.lower() in resume_text for k in kws) for sec, kws in ats_sections.items()}

        # Step 4: Semantic similarity
        resume_sents = sent_tokenize(resume_text)
        jd_sents = sent_tokenize(job_desc)

        resume_embeddings = model.encode(resume_sents, convert_to_tensor=True)
        jd_embeddings = model.encode(jd_sents, convert_to_tensor=True)

        cos_scores = util.pytorch_cos_sim(jd_embeddings, resume_embeddings)
        max_sim_per_jd = cos_scores.max(dim=1).values
        semantic_similarity = max_sim_per_jd.mean().item()

        # Step 5: Compute scores
        section_score = sum(sections_detected.values()) * 5
        keyword_score = len(matched) / max(len(job_tokens), 1) * 35
        semantic_score = semantic_similarity * 60
        if semantic_similarity < 0.6:
            semantic_score *= 0.5

        ats_score = round(min(100, section_score + keyword_score + semantic_score), 2)

        # Step 6: Suggestions
        suggestions = []
        job_keywords = list(job_tokens - resume_tokens)
        if job_keywords:
            job_kw_embeddings = model.encode(job_keywords, convert_to_tensor=True)
            resume_embedding = model.encode(resume_text, convert_to_tensor=True)
            sims = util.cos_sim(job_kw_embeddings, resume_embedding).squeeze()

            ranked_keywords = sorted(zip(job_keywords, sims.tolist()), key=lambda x: x[1])
            ignore = {
                "high", "good", "able", "using", "based", "help", "work", "team",
                "time", "will", "you", "your", "well", "effort", "performing",
                "leading", "cross", "reviewing", "responsible", "ensure", "support"
            }

            filtered_keywords = []
            for kw, score in ranked_keywords:
                if kw in ignore or kw in stopwords.words('english'):
                    continue
                pos = pos_tag([kw])[0][1]
                if pos.startswith(('NN', 'VB', 'JJ')):
                    filtered_keywords.append(kw)

            top_keywords = filtered_keywords[:17]
            if top_keywords:
                suggestions.append(
                    f"Consider adding these relevant keywords to better align with the job description: {', '.join(top_keywords)}."
                )
            else:
                suggestions.append("Your resume semantically covers most of the job description well.")
        else:
            suggestions.append("No missing keywords detected.")

        # ✅ Add this resume’s result
        results.append({
            "filename": resume.filename,
            "ats_score": ats_score,
            "sections_detected": sections_detected,
            "semantic_similarity": round(semantic_similarity * 100, 2),
            "keyword_match": {
                "match_percent": round(len(matched) / max(len(job_tokens), 1) * 100, 2),
                "matched_keywords": sorted(matched),
            },
            "improvements": suggestions,
        })

    # ✅ Final Return
    if len(results) == 1:
        return results[0]
    else:
        return {"results": results}

def generate_response(message: str, system_prompt: str, temperature: float, max_tokens: int):
    conversation = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": message}
    ]

    response = client.chat.completions.create(
        model="llama-3.1-8B-Instant",
        messages=conversation,
        temperature=temperature,
        max_tokens=max_tokens,
        stream=False
    )

    return response.choices[0].message.content

@app.post("/analyze-resume")
async def analyze_resume_endpoint(
    resume_text: str = Form(...),
    job_description: str = Form(...),
    with_job_description: bool = Form(True),
    temperature: float = Form(0.3),
    max_tokens: int = Form(1500)
):
    

    if with_job_description and job_description.strip():
           prompt = f"""
You are an expert ATS (Applicant Tracking System) evaluator.
Analyze the following resume in the context of the provided job description.
Your analysis should cover:
- Match percentage (semantic and keyword-based)
- Missing or weak keywords
- Summary 
- Actionable recommendations for improvement 

Job Description:
{job_description}

Resume:
{resume_text}
"""
    else:
        prompt = f"""
You are an expert ATS resume evaluator.
Analyze the following resume without a specific job description.
Your analysis should cover:
- Overall resume quality score (0–10)
- Evaluation based on Impact, Clarity, Structure, and Skills Relevance
- 2–3 line summary
- 3–4 actionable improvement points

Resume:
{resume_text}
"""

    # ai_output = generate_response(prompt, "You are an expert ATS resume analyzer.", temperature, max_tokens)

    # return JSONResponse(content={"ai_analysis": ai_output})
    ai_output = generate_response(
        message=prompt,
        system_prompt="You are an expert ATS resume analyzer. Respond only in plain text, without section headers or markdown markers.",
        temperature=temperature,
        max_tokens=max_tokens
    )

    # Step 4: Return structured JSON response
    return JSONResponse(content={
        "ai_analysis": ai_output,
        "input_summary": {
            "with_job_description": with_job_description,
            "job_description_present": bool(job_description.strip()),
        }
    })
